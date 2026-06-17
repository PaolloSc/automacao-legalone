"""
Extrator de dados de petições via Azure Document Intelligence + Groq (Llama).
Opção B: OCR/extração de texto de PDF/DOCX/DOC → Groq extrai campos estruturados (grátis).
"""

import io
import json
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
load_dotenv()

from config_azure import (
    AZURE_DOC_INTELLIGENCE_CONFIG,
    CAMPOS_OBRIGATORIOS,
    TIPOS_CADASTRO_VALIDOS,
)

logger = logging.getLogger(__name__)

GROQ_API_KEY = os.getenv('GROQ_API_KEY', '')
GROQ_MODEL = os.getenv('GROQ_MODEL', 'llama-3.3-70b-versatile')


class PeticaoExtractor:
    """Extrai campos estruturados de petições judiciais usando Azure OCR + Groq."""

    def __init__(self):
        self._validate_config()

    def _validate_config(self):
        if not AZURE_DOC_INTELLIGENCE_CONFIG['endpoint']:
            logger.warning("AZURE_DOC_INTELLIGENCE_ENDPOINT não configurado")
        if not GROQ_API_KEY:
            logger.warning("GROQ_API_KEY não configurada")

    def ocr_pdf(self, pdf_path: str) -> str:
        """Extrai texto do PDF via Azure Document Intelligence."""
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential

        client = DocumentIntelligenceClient(
            endpoint=AZURE_DOC_INTELLIGENCE_CONFIG['endpoint'],
            credential=AzureKeyCredential(AZURE_DOC_INTELLIGENCE_CONFIG['api_key']),
        )

        with open(pdf_path, "rb") as f:
            poller = client.begin_analyze_document("prebuilt-read", body=f)

        result = poller.result()
        texto = result.content
        logger.info(f"OCR extraiu {len(texto)} caracteres de {Path(pdf_path).name}")
        return texto

    def ocr_bytes(self, pdf_bytes: bytes) -> str:
        """Extrai texto de bytes do PDF via Azure Document Intelligence."""
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential

        client = DocumentIntelligenceClient(
            endpoint=AZURE_DOC_INTELLIGENCE_CONFIG['endpoint'],
            credential=AzureKeyCredential(AZURE_DOC_INTELLIGENCE_CONFIG['api_key']),
        )

        poller = client.begin_analyze_document("prebuilt-read", body=pdf_bytes)
        result = poller.result()
        texto = result.content
        logger.info(f"OCR extraiu {len(texto)} caracteres do PDF em memória")
        return texto

    def extrair_campos(self, texto: str) -> dict:
        """Usa Groq (Llama 3.3 70B) para extrair campos estruturados — grátis."""
        from groq import Groq

        client = Groq(api_key=GROQ_API_KEY)

        prompt_sistema = """Você é um assistente jurídico especializado em extrair dados de petições judiciais brasileiras de qualquer área do direito (cível, trabalhista, tributário, previdenciário, criminal, família, consumidor, etc.).

Extraia os seguintes campos do texto da petição. Retorne APENAS um JSON válido, sem markdown, sem explicações:

{
  "cnj": "número do processo no formato NNNNNNN-NN.NNNN.N.NN.NNNN",
  "cliente": "nome da parte representada pelo escritório Carvalho & Furtado",
  "contrario": "nome da parte adversa",
  "natureza": "tipo da ação (ex: Reclamação Trabalhista, Ação de Cobrança, Mandado de Segurança)",
  "tribunal": "sigla do tribunal (ex: TJMG, TRT3, TST, TRF1, STJ)",
  "comarca": "cidade/localidade da vara",
  "instancia": "1ª Instância, 2ª Instância ou Instância Superior",
  "posicao": "posição do cliente: Autor, Réu, Reclamante, Reclamado, etc.",
  "fase": "Conhecimento, Recursal, Execução ou Cumprimento de Sentença",
  "tipo_cadastro": "CADASTRO INICIAL, DECISÕES, RECURSO, ARQUIVAMENTO COMPLETO ou ARQUIVAMENTO SIMPLES"
}

Regras:
- Se um campo não for encontrado, use "NAO LOCALIZADO"
- Não invente dados
- Valide consistência (TRT = trabalhista, TJMG = cível em MG)
- tipo_cadastro deve ser um dos 5 valores exatos listados acima"""

        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": f"Extraia os campos desta petição:\n\n{texto[:8000]}"},
            ],
            temperature=0.1,
            max_tokens=1000,
        )

        resposta = response.choices[0].message.content.strip()

        if resposta.startswith("```"):
            resposta = resposta.split("```")[1]
            if resposta.startswith("json"):
                resposta = resposta[4:]
            resposta = resposta.strip()

        campos = json.loads(resposta)
        self._validar_campos(campos)
        logger.info(f"Campos extraídos: CNJ={campos.get('cnj', '?')}, Cliente={campos.get('cliente', '?')}")
        return campos

    def _validar_campos(self, campos: dict):
        """Valida e normaliza campos extraídos."""
        for campo in CAMPOS_OBRIGATORIOS:
            if campo not in campos:
                campos[campo] = "NAO LOCALIZADO"

        tipo = campos.get('tipo_cadastro', '')
        if tipo and tipo.upper() not in [t.upper() for t in TIPOS_CADASTRO_VALIDOS]:
            logger.warning(f"tipo_cadastro '{tipo}' inválido, usando CADASTRO INICIAL")
            campos['tipo_cadastro'] = 'CADASTRO INICIAL'

    # ------------------------------------------------------------------
    # Extração de texto por tipo de arquivo
    # ------------------------------------------------------------------
    def extrair_texto_docx(self, caminho_ou_bytes) -> str:
        """Extrai texto de DOCX via python-docx."""
        from docx import Document

        if isinstance(caminho_ou_bytes, bytes):
            doc = Document(io.BytesIO(caminho_ou_bytes))
        else:
            doc = Document(caminho_ou_bytes)

        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        # Também extrai texto de tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                if celulas:
                    paragrafos.append(' | '.join(celulas))

        texto = '\n'.join(paragrafos)
        logger.info(f"DOCX extraiu {len(texto)} caracteres")
        return texto

    def extrair_texto_doc(self, caminho_ou_bytes) -> str:
        """Extrai texto de DOC (formato antigo) via antiword ou fallback Document Intelligence."""
        if isinstance(caminho_ou_bytes, bytes):
            # Salva em temp pra processar
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(caminho_ou_bytes)
                tmp_path = tmp.name
        else:
            tmp_path = str(caminho_ou_bytes)

        try:
            # Tenta antiword (se instalado)
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                texto = result.stdout.strip()
                logger.info(f"DOC (antiword) extraiu {len(texto)} caracteres")
                return texto
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Fallback: envia pro Document Intelligence (aceita DOC)
        logger.info("[DOC] antiword indisponível, usando Document Intelligence como fallback")
        if isinstance(caminho_ou_bytes, bytes):
            return self.ocr_bytes(caminho_ou_bytes)
        else:
            return self.ocr_pdf(tmp_path)
        finally:
            if isinstance(caminho_ou_bytes, bytes):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

    def _detectar_tipo(self, nome_arquivo: str) -> str:
        """Detecta tipo do arquivo pela extensão."""
        ext = Path(nome_arquivo).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.doc':
            return 'doc'
        else:
            return 'desconhecido'

    def extrair_texto_arquivo(self, caminho_ou_bytes, nome_arquivo: str = '') -> str:
        """Extrai texto de qualquer formato suportado (PDF, DOCX, DOC)."""
        tipo = self._detectar_tipo(nome_arquivo)

        if tipo == 'pdf':
            if isinstance(caminho_ou_bytes, bytes):
                return self.ocr_bytes(caminho_ou_bytes)
            return self.ocr_pdf(caminho_ou_bytes)
        elif tipo == 'docx':
            return self.extrair_texto_docx(caminho_ou_bytes)
        elif tipo == 'doc':
            return self.extrair_texto_doc(caminho_ou_bytes)
        else:
            # Tenta Document Intelligence como fallback genérico
            logger.warning(f"Tipo '{tipo}' desconhecido para '{nome_arquivo}', tentando Document Intelligence")
            if isinstance(caminho_ou_bytes, bytes):
                return self.ocr_bytes(caminho_ou_bytes)
            return self.ocr_pdf(caminho_ou_bytes)

    # ------------------------------------------------------------------
    # Pipelines completos
    # ------------------------------------------------------------------
    def processar_pdf(self, pdf_path: str) -> dict:
        """Pipeline completo: OCR → extração de campos."""
        texto = self.ocr_pdf(pdf_path)
        return self.extrair_campos(texto)

    def processar_pdf_bytes(self, pdf_bytes: bytes) -> dict:
        """Pipeline completo a partir de bytes: OCR → extração."""
        texto = self.ocr_bytes(pdf_bytes)
        return self.extrair_campos(texto)

    def processar_arquivo(self, caminho_ou_bytes, nome_arquivo: str = '') -> dict:
        """Pipeline completo para qualquer formato: PDF, DOCX, DOC."""
        texto = self.extrair_texto_arquivo(caminho_ou_bytes, nome_arquivo)
        return self.extrair_campos(texto)

    def processar_arquivo_bytes(self, arquivo_bytes: bytes, nome_arquivo: str) -> dict:
        """Pipeline completo a partir de bytes com detecção automática de formato."""
        texto = self.extrair_texto_arquivo(arquivo_bytes, nome_arquivo)
        return self.extrair_campos(texto)
